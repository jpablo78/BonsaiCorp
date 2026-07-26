"""Genera el informe ejecutivo y el anexo técnico de Bonsai Corp.

Los documentos se construyen a partir de resultados ya validados de la
solución validada y de sensibilidades que se etiquetan expresamente como
diagnósticas. No reejecuta la optimización.
"""

from __future__ import annotations

import os
from collections import defaultdict
from pathlib import Path

os.environ.setdefault("MPLBACKEND", "Agg")
os.environ.setdefault("MPLCONFIGDIR", str(Path(__file__).resolve().parents[1] / ".mplconfig"))

import matplotlib.pyplot as plt
import pandas as pd
from matplotlib.ticker import FuncFormatter
from docxcompose.composer import Composer
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from bonsai.config import DISCOUNT_TIERS, FreightPolicy, PALLET_LENGTH_MM, PALLET_MAX_HEIGHT_MM, PALLET_WIDTH_MM
from bonsai.costs import box_type_key, tier_index
from bonsai.data import load_prepared_data
from bonsai.decimal_io import validate_decimal_solution_csv
from bonsai.geometry import boxes_per_pallet


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output_documentos"
OUT.mkdir(exist_ok=True)

NAVY = "17324D"
TEAL = "007F86"
GREEN = "2E7D32"
LIGHT = "EAF1F5"
GRAY = "5B6573"

HISTORICAL_TOTAL = 209_235_093.6821712
WINNING_TOTAL = 188_079_384.24
SAVINGS = HISTORICAL_TOTAL - WINNING_TOTAL


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shade = OxmlElement("w:shd")
    shade.set(qn("w:fill"), fill)
    tc_pr.append(shade)


def set_cell_text(cell, text: str, *, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(9)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_number(section) -> None:
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Bonsai Corp | ")
    run.font.size = Pt(8)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    p._p.append(field)


def base_doc(title: str, subtitle: str) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    add_page_number(section)
    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10)
    for style_name, size, color in [("Title", 27, NAVY), ("Heading 1", 17, NAVY), ("Heading 2", 12, TEAL)]:
        style = styles[style_name]
        style.font.name = "Aptos Display" if style_name != "Normal" else "Aptos"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("ALIXPARTNERS DATA CHALLENGE | BONSAI CORP")
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string(TEAL)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(27)
    r.font.color.rgb = RGBColor.from_string(NAVY)
    p = doc.add_paragraph(subtitle)
    p.paragraph_format.space_after = Pt(16)
    p.runs[0].font.color.rgb = RGBColor.from_string(GRAY)
    return doc


def add_source(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Fuente: " + text)
    r.italic = True
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor.from_string(GRAY)


def add_callout(doc: Document, title: str, body: str, fill: str = "E8F3F1") -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(7)
    r = p.add_run(title + "\n")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(NAVY)
    r = p.add_run(body)
    r.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(7)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, header in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, NAVY)
        set_cell_text(cell, header, bold=True, color="FFFFFF")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            if len(table.rows) % 2 == 1:
                set_cell_shading(cells[i], "F4F7F9")
    if widths:
        for row in table.rows:
            for cell, width in zip(row.cells, widths):
                cell.width = Cm(width)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def make_chart_costs() -> Path:
    path = OUT / "grafico_costos.png"
    names = ["Histórico", "Solución\noptimizada"]
    values = [HISTORICAL_TOTAL / 1e6, WINNING_TOTAL / 1e6]
    fig, ax = plt.subplots(figsize=(6.7, 3.35))
    bars = ax.bar(names, values, color=["#AAB8C2", "#007F86"], width=0.55)
    ax.set_ylabel("USD millones")
    ax.set_ylim(0, 240)
    ax.spines[["top", "right"]].set_visible(False)
    ax.grid(axis="y", alpha=0.18)
    for bar, value in zip(bars, values):
        ax.text(bar.get_x() + bar.get_width() / 2, value + 4, f"{value:,.1f}", ha="center", fontweight="bold")
    y_bracket = 218
    # Llave entre los bordes interiores de las barras: no debe cubrir las etiquetas.
    ax.plot([0.30, 0.30, 0.70, 0.70], [values[0], y_bracket, y_bracket, values[1]], color="#2E7D32", linestyle="--", linewidth=1.5)
    ax.text(0.5, y_bracket + 4, f"−USD {SAVINGS / 1e6:,.1f} M\n−10,1%", ha="center", color="#2E7D32", fontweight="bold")
    fig.tight_layout()
    fig.savefig(path, dpi=180, bbox_inches="tight")
    plt.close(fig)
    return path


def make_chart_thickness() -> Path:
    path = OUT / "grafico_componentes_costo.png"
    names = ["Histórico", "Solución\noptimizada"]
    packaging = [30.166293682, 26.92118424]
    freight = [179.0688, 161.1582]
    fig, ax = plt.subplots(figsize=(6.7, 3.2))
    ax.bar(names, freight, color="#AAB8C2", width=0.58, label="Flete")
    ax.bar(names, packaging, bottom=freight, color="#007F86", width=0.58, label="Cartón")
    ax.set_ylabel("USD millones")
    ax.set_ylim(0, 230)
    ax.spines[["top", "right"]].set_visible(False)
    ax.grid(axis="y", alpha=0.18)
    for i, total in enumerate([sum(v) for v in zip(packaging, freight)]):
        ax.text(i, total + 5, f"{total:,.1f}", ha="center", fontweight="bold", fontsize=9)
    # La leyenda va por fuera del área de dibujo: nunca debe tapar las barras.
    ax.legend(frameon=False, loc="upper center", bbox_to_anchor=(0.5, 1.18), ncol=2)
    fig.subplots_adjust(top=0.78, bottom=0.16, left=0.12, right=0.98)
    fig.savefig(path, dpi=180, bbox_inches="tight")
    plt.close(fig)
    return path


def make_chart_operating_scale() -> Path:
    path = OUT / "grafico_operacion.png"
    fig, axes = plt.subplots(1, 2, figsize=(6.7, 2.8))
    for ax, labels, values, title in [
        (axes[0], ["Actual", "Optimizado"], [204, 59], "Tipos de caja"),
        (axes[1], ["Actual", "Optimizado"], [1_193_792, 1_074_388], "Pallets anuales"),
    ]:
        bars = ax.bar(labels, values, color=["#AAB8C2", "#007F86"], width=0.55)
        ax.set_title(title, fontsize=11, fontweight="bold", color="#17324D")
        ax.spines[["top", "right"]].set_visible(False)
        ax.grid(axis="y", alpha=0.18)
        for bar, value in zip(bars, values):
            label = f"{value:,}".replace(",", ".")
            ax.text(bar.get_x() + bar.get_width() / 2, value * 1.03, label, ha="center", fontweight="bold", fontsize=8)
    axes[0].set_ylim(0, 240)
    axes[1].set_ylim(0, 1_350_000)
    axes[1].set_ylabel("Millones de pallets", fontsize=8)
    axes[1].yaxis.set_major_formatter(
        FuncFormatter(lambda value, _position: f"{value / 1_000_000:.1f}".replace(".", ","))
    )
    fig.tight_layout()
    fig.savefig(path, dpi=180, bbox_inches="tight")
    plt.close(fig)
    return path


def _eda_frames() -> tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame]:
    spec = pd.read_csv(ROOT / "especificaciones_cajas.csv", encoding="utf-8-sig")
    ops = pd.read_csv(ROOT / "operaciones_planta.csv", encoding="utf-8-sig")
    procurement = pd.read_csv(ROOT / "procurement_cajas.csv", encoding="utf-8-sig")
    return spec, ops, procurement


def solution_metrics() -> dict[str, object]:
    """Calcula indicadores comparables del catálogo actual y de la solución final."""
    data = load_prepared_data(ROOT)
    final = validate_decimal_solution_csv(
        ROOT / "baseline" / "asignacion_0_1mm.csv",
        data,
        FreightPolicy(extra_region_share=0.0),
        required_thickness_mm=3.0,
    )
    current = {p.code: data.current_boxes[p.current_box_type_id] for p in data.products}

    def summarize(assignment, is_final: bool) -> dict[str, object]:
        total_demand = sum(p.annual_volume for p in data.products)
        u_caja = sum(
            p.annual_volume * p.product_volume_mm3 / assignment[p.code].internal.volume_mm3
            for p in data.products
        ) / total_demand
        occupied = defaultdict(float)
        capacity = defaultdict(float)
        pallets = defaultdict(int)
        tier_volumes = defaultdict(int)
        for p in data.products:
            box = assignment[p.code]
            per_pallet = box.capacity_per_pallet if is_final else boxes_per_pallet(box.external)
            key = box_type_key(box)
            for plant, volume in p.annual_volume_by_plant.items():
                if not volume:
                    continue
                count = (volume + per_pallet - 1) // per_pallet
                pallets[plant] += count
                occupied[plant] += volume * box.external.volume_mm3
                capacity[plant] += count * PALLET_LENGTH_MM * PALLET_WIDTH_MM * PALLET_MAX_HEIGHT_MM
                tier_volumes[key, plant] += volume
        tier_counts = [0] * len(DISCOUNT_TIERS)
        tier_units = [0] * len(DISCOUNT_TIERS)
        for volume in tier_volumes.values():
            index = tier_index(volume)
            tier_counts[index] += 1
            tier_units[index] += volume
        u_pallet = {plant: occupied[plant] / capacity[plant] for plant in pallets}
        return {
            "u_caja": u_caja,
            "pallets": dict(pallets),
            "u_pallet": u_pallet,
            "u_pallet_total": sum(occupied.values()) / sum(capacity.values()),
            "tier_counts": tier_counts,
            "tier_units": tier_units,
        }

    before = summarize(current, False)
    after = summarize(final.assignment, True)
    capacity_change = {"better": 0, "same": 0, "worse": 0}
    for product in data.products:
        old = boxes_per_pallet(current[product.code].external)
        new = final.assignment[product.code].capacity_per_pallet
        capacity_change["better" if new > old else "worse" if new < old else "same"] += 1
    return {"before": before, "after": after, "capacity_change": capacity_change}


def make_chart_dimensions() -> Path:
    path = OUT / "eda_distribucion_dimensiones.png"
    spec, _, _ = _eda_frames()
    fields = [
        ("caja_interior_largo", "Largo interno"),
        ("caja_interior_ancho", "Ancho interno"),
        ("caja_interior_alto", "Alto interno"),
    ]
    fig, axes = plt.subplots(1, 3, figsize=(7.1, 2.65))
    for ax, (field, label) in zip(axes, fields):
        ax.hist(spec[field], bins=12, color="#007F86", edgecolor="white")
        ax.axvline(spec[field].median(), color="#17324D", linewidth=1.4, linestyle="--", label="Mediana")
        ax.set_title(label, fontsize=10, fontweight="bold", color="#17324D")
        ax.set_xlabel("mm")
        ax.spines[["top", "right"]].set_visible(False)
        ax.grid(axis="y", alpha=0.15)
    axes[0].set_ylabel("Tipos de caja")
    axes[2].legend(frameon=False, fontsize=8, loc="upper right")
    fig.tight_layout()
    fig.savefig(path, dpi=180, bbox_inches="tight")
    plt.close(fig)
    return path


def make_chart_utilization() -> Path:
    path = OUT / "eda_utilizacion_pallet.png"
    spec, ops, _ = _eda_frames()
    catalog = pd.read_csv(ROOT / "catalogo_productos.csv", encoding="utf-8-sig")
    demand = catalog[["codigo_producto", "caja_tipo_id"]].merge(ops[["codigo_producto", "volumen_producto_total"]], on="codigo_producto").groupby("caja_tipo_id")["volumen_producto_total"].sum().rename("volumen_total")
    combined = spec.merge(demand, on="caja_tipo_id", how="left").fillna({"volumen_total": 0})
    weighted = (combined["utilizacion"] * combined["volumen_total"]).sum() / combined["volumen_total"].sum()
    fig, ax = plt.subplots(figsize=(6.8, 2.85))
    ax.hist(spec["utilizacion"] * 100, bins=12, color="#AAB8C2", edgecolor="white")
    ax.axvline(spec["utilizacion"].mean() * 100, color="#17324D", linestyle="--", linewidth=1.5, label="Promedio simple: 83,7%")
    ax.axvline(weighted * 100, color="#007F86", linewidth=1.8, label=f"Ponderado por demanda: {weighted * 100:.1f}%")
    ax.set_xlabel("Utilización de pallet (%)")
    ax.set_ylabel("Tipos de caja")
    ax.spines[["top", "right"]].set_visible(False)
    ax.grid(axis="y", alpha=0.15)
    ax.legend(frameon=False, loc="upper left", fontsize=8)
    fig.tight_layout()
    fig.savefig(path, dpi=180, bbox_inches="tight")
    plt.close(fig)
    return path


def make_chart_plants() -> Path:
    path = OUT / "eda_plantas.png"
    _, ops, _ = _eda_frames()
    plant_names = ["Buenos Aires", "Curitiba", "Santiago", "Monterrey", "Bakersfield"]
    suffixes = ["buenos_aires", "curitiba", "santiago", "monterrey", "bakersfield"]
    demand = [ops[f"volumen_producto_planta_{s}"].sum() / 1e6 for s in suffixes]
    pallets = [ops[f"cantidad_pallets_planta_{s}"].sum() / 1e3 for s in suffixes]
    fig, axes = plt.subplots(1, 2, figsize=(7.1, 2.85))
    axes[0].barh(plant_names, demand, color="#007F86")
    axes[0].set_title("Demanda anual", fontsize=10, fontweight="bold", color="#17324D")
    axes[0].set_xlabel("Millones de unidades")
    axes[1].barh(plant_names, pallets, color="#AAB8C2")
    axes[1].set_title("Pallets actuales", fontsize=10, fontweight="bold", color="#17324D")
    axes[1].set_xlabel("Miles de pallets")
    for ax in axes:
        ax.spines[["top", "right"]].set_visible(False)
        ax.grid(axis="x", alpha=0.15)
    axes[1].set_yticks([])
    fig.tight_layout()
    fig.savefig(path, dpi=180, bbox_inches="tight")
    plt.close(fig)
    return path


def make_chart_pareto() -> Path:
    path = OUT / "eda_pareto_tipos.png"
    _, ops, _ = _eda_frames()
    catalog = pd.read_csv(ROOT / "catalogo_productos.csv", encoding="utf-8-sig")
    volumes = (
        catalog[["codigo_producto", "caja_tipo_id"]]
        .merge(ops[["codigo_producto", "volumen_producto_total"]], on="codigo_producto")
        .groupby("caja_tipo_id")["volumen_producto_total"].sum()
        .sort_values(ascending=False).reset_index(drop=True)
    )
    cumulative = volumes.cumsum() / volumes.sum() * 100
    fig, ax = plt.subplots(figsize=(6.8, 2.85))
    ax.bar(range(1, len(volumes) + 1), volumes / 1e6, color="#AAB8C2", width=1.0)
    ax.set_xlabel("Tipos de caja actuales, ordenados por volumen")
    ax.set_ylabel("Volumen anual (millones)")
    ax.spines[["top", "right"]].set_visible(False)
    ax.grid(axis="y", alpha=0.15)
    ax2 = ax.twinx()
    ax2.plot(range(1, len(volumes) + 1), cumulative, color="#007F86", linewidth=2)
    ax2.set_ylabel("Volumen acumulado (%)", color="#007F86")
    ax2.set_ylim(0, 105)
    ax2.axhline(80, color="#17324D", linestyle="--", linewidth=1)
    fig.tight_layout()
    fig.savefig(path, dpi=180, bbox_inches="tight")
    plt.close(fig)
    return path


def make_chart_utilization_before_after() -> Path:
    path = OUT / "comparacion_utilizacion_pallet.png"
    metrics = solution_metrics()
    plants = ["Buenos Aires", "Curitiba", "Santiago", "Monterrey", "Bakersfield"]
    keys = ["buenos_aires", "curitiba", "santiago", "monterrey", "bakersfield"]
    before = [metrics["before"]["u_pallet"][key] * 100 for key in keys]
    after = [metrics["after"]["u_pallet"][key] * 100 for key in keys]
    fig, ax = plt.subplots(figsize=(6.8, 3.45))
    fig.suptitle(
        "Utilización promedio del pallet por región y escenario",
        fontsize=12,
        fontweight="bold",
        color="#17324D",
        y=0.99,
    )
    x = list(range(len(plants)))
    ax.bar([v - 0.19 for v in x], before, width=0.38, label="Actual", color="#AAB8C2")
    ax.bar([v + 0.19 for v in x], after, width=0.38, label="Optimizado", color="#007F86")
    ax.set_xticks(x, plants, rotation=18, ha="right")
    ax.set_ylim(0, 100)
    ax.set_ylabel("Utilización volumétrica de pallet (%)")
    ax.grid(axis="y", alpha=0.15)
    ax.spines[["top", "right"]].set_visible(False)
    ax.legend(
        frameon=False,
        ncol=2,
        loc="upper center",
        bbox_to_anchor=(0.5, 1.16),
        borderaxespad=0.0,
    )
    fig.subplots_adjust(left=0.12, right=0.98, bottom=0.24, top=0.80)
    fig.savefig(path, dpi=180, bbox_inches="tight")
    plt.close(fig)
    return path


def make_chart_tradeoff() -> Path:
    path = OUT / "tradeoff_utilizacion.png"
    metrics = solution_metrics()
    fig, axes = plt.subplots(1, 2, figsize=(6.8, 2.85))
    internal_use = [metrics["before"]["u_caja"] * 100, metrics["after"]["u_caja"] * 100]
    bars = axes[0].bar(["Actual", "Optimizado"], internal_use, color=["#AAB8C2", "#007F86"])
    axes[0].set_title("Aprovechamiento interno", fontweight="bold", color="#17324D")
    axes[0].set_ylim(0, 100)
    axes[0].set_ylabel("Uso del volumen interno (%)", fontsize=8)
    for bar, value in zip(bars, internal_use):
        axes[0].text(
            bar.get_x() + bar.get_width() / 2,
            value - 4.0,
            f"{value:.1f}%".replace(".", ","),
            ha="center",
            va="top",
            fontweight="bold",
            fontsize=8,
            color="#17324D",
        )
    changes = metrics["capacity_change"]
    values = [changes["better"], changes["same"], changes["worse"]]
    bars = axes[1].bar(["Aumenta", "Sin cambio", "Disminuye"], values, color=["#007F86", "#AAB8C2", "#C8D3DA"])
    axes[1].set_title("Capacidad de pallet por SKU", fontweight="bold", color="#17324D")
    axes[1].set_ylabel("Cantidad de SKU")
    axes[1].set_ylim(0, 290)
    for bar, value in zip(bars, values):
        axes[1].text(bar.get_x() + bar.get_width() / 2, value + 7, str(value), ha="center", fontweight="bold", fontsize=8)
    for ax in axes:
        ax.spines[["top", "right"]].set_visible(False)
        ax.grid(axis="y", alpha=0.15)
    fig.tight_layout()
    fig.savefig(path, dpi=180, bbox_inches="tight")
    plt.close(fig)
    return path


def make_chart_tiers() -> Path:
    path = OUT / "tiers_antes_despues.png"
    metrics = solution_metrics()
    names = ["T1", "T2", "T3", "T4", "T5"]
    before = metrics["before"]["tier_units"]
    after = metrics["after"]["tier_units"]
    fig, ax = plt.subplots(figsize=(6.8, 3.0))
    x = list(range(len(names)))
    ax.bar([v - 0.19 for v in x], [v / 1e6 for v in before], width=0.38, label="Actual", color="#AAB8C2")
    ax.bar([v + 0.19 for v in x], [v / 1e6 for v in after], width=0.38, label="Optimizado", color="#007F86")
    ax.set_xticks(x, names)
    ax.set_ylabel("Volumen anual (millones de unidades)")
    ax.set_title("Volumen por banda de Procurement (tipo–planta)", fontweight="bold", color="#17324D")
    ax.legend(frameon=False, ncol=2)
    ax.grid(axis="y", alpha=0.15)
    ax.spines[["top", "right"]].set_visible(False)
    fig.tight_layout()
    fig.savefig(path, dpi=180, bbox_inches="tight")
    plt.close(fig)
    return path


def executive_report() -> Path:
    doc = base_doc(
        "Packaging: una plataforma más simple, con USD 21,2 M de ahorro anual",
        "Recomendación ejecutiva para Bonsai Corp | Juan Pablo Villaverde | Julio de 2026",
    )
    add_callout(
        doc,
        "Decisión propuesta",
        "Adoptar el catálogo optimizado de 59 tipos de caja, con un único espesor global de 3,0 mm. La solución reduce el costo anual estimado de USD 209,2 M a USD 188,1 M, generando un ahorro de USD 21,2 M (10,11%).",
    )
    doc.add_heading("Conclusión ejecutiva", level=1)
    doc.add_paragraph(
        "Bonsai puede lograr una reducción significativa y recurrente de costos, manteniendo un único espesor y sin introducir excepciones por planta. La recomendación concentra la complejidad donde realmente agrega valor: el diseño dimensional y la consolidación del portafolio. Con un espesor global de 3,0 mm, el catálogo operativo se reduce de 204 tipos de caja a 59, simplificando la operación y reduciendo costos de forma recurrente."
    )
    doc.add_paragraph(
        "Además, se entrega una solución escalable y reutilizable: permite actualizar el análisis con nuevas proyecciones de demanda, asignar nuevos SKU a cajas existentes y reoptimizar el portafolio cuando cambien las condiciones relevantes."
    )
    doc.add_picture(str(make_chart_costs()), width=Inches(6.45))

    doc.add_heading("Qué se logró", level=1)
    add_table(doc, ["Indicador", "Resultado", "Implicancia"], [
        ["Costo anual total", "USD 188,1 M", "USD 21,2 M menos que el histórico"],
        ["Ahorro estimado", "10,11%", "Resultado verificado por el evaluador oficial"],
        ["Tipos de caja", "59", "−145 tipos frente a los 204 actuales"],
        ["Espesor", "3,0 mm global", "Un estándar único y factible"],
        ["Pallets anuales", "1.074.388", "Costo logístico calculado a USD 150 por pallet"],
    ], [4.0, 3.0, 9.0])

    doc.add_heading("Punto de partida: un portafolio fragmentado", level=1)
    doc.add_paragraph(
        "El análisis exploratorio identifica 427 productos atendidos por 204 tipos de caja actuales en cinco plantas. Las cajas actuales presentan una utilización promedio simple de pallet de 83,7% y una utilización ponderada por demanda de 83,2%. La oportunidad no era sólo reducir cartón: existía dispersión de diseños y capacidad de pallet no aprovechada de forma uniforme."
    )
    add_table(doc, ["Dimensión interna actual", "Mínimo", "Mediana", "Máximo"], [
        ["Largo", "375 mm", "386 mm", "395 mm"],
        ["Ancho", "182 mm", "286 mm", "297 mm"],
        ["Alto", "138 mm", "222 mm", "363 mm"],
        ["Utilización de pallet", "64,0%", "84,0%", "99,0%"],
    ], [5.0, 3.5, 3.5, 3.5])

    doc.add_heading("Por qué esta solución es sólida", level=1)
    doc.add_paragraph(
        "La solución no es una reducción lineal de medidas. Para cada producto se evaluaron configuraciones dimensionales que respetan tolerancias, capacidad, headspace, resistencia ECT y palletización. Después se eligió el conjunto de diseños que minimiza conjuntamente costo de cartón y pallets, y que permite que productos compatibles compartan una misma caja."
    )
    doc.add_paragraph(
        "La geometría se expresa con precisión de 0,1 mm. El evaluador oficial aceptó ese nivel de precisión y el resultado obtenido con esa especificación fue consistente con el cálculo local. En producción, las cotas deberán traducirse a tolerancias industriales y validarse con Ingeniería de Packaging antes de aprobar herramentales o emitir órdenes de compra.")
    doc.add_paragraph(
        "Como control adicional, el MIP de certificación alcanzó optimalidad con gap de 0,0000% en su universo discreto de diseños a 0,1 mm. En consecuencia, no existe una alternativa de menor costo entre los 1.336 diseños candidatos de esa corrida. Esta garantía se limita al universo discreto generado por el premodelado; no afirma optimalidad sobre dimensiones continuas no representadas.")

    doc.add_heading("De dónde proviene el ahorro", level=1)
    doc.add_picture(str(make_chart_thickness()), width=Inches(6.45))
    doc.add_paragraph("El ahorro se reparte entre USD 3,25 M de cartón y USD 17,91 M de flete. La mejora de flete proviene de 119.404 pallets menos por año; el modelo monetiza ambos componentes de forma directa, en lugar de optimizarlos de manera separada.")
    doc.add_picture(str(make_chart_operating_scale()), width=Inches(6.45))
    doc.add_paragraph("Los costos históricos entregados son consistentes con una tarifa de USD 150 por pallet: las 427 filas cumplen exactamente la relación pallets × USD 150 y los datos no contienen volumen ni fletes extra-zona.")

    doc.add_heading("El ajuste físico preserva la eficiencia logística", level=1)
    doc.add_picture(str(make_chart_tradeoff()), width=Inches(6.45))
    doc.add_paragraph("La consolidación utiliza, en promedio ponderado por demanda, 96,7% del volumen interno de cada caja, frente al 100,0% de la referencia. Esa diferencia permite adaptar las geometrías y compartir diseños, sin reducir la capacidad de pallet de ningún SKU: aumenta para 275 productos y se mantiene para 152. A nivel de la red, la utilización total de pallet mejora de 82,2% a 93,5%, habilitando 119.404 pallets anuales menos.")

    doc.add_heading("Consideraciones para la implementación", level=1)
    for item in [
        "Validar físicamente los diseños mediante pruebas de ajuste, resistencia y apilado.",
        "Traducir las dimensiones del modelo a especificaciones industriales con tolerancias de fabricación definidas.",
        "Realizar un piloto controlado antes de reemplazar el catálogo vigente.",
        "Planificar una transición controlada que garantice continuidad de abastecimiento.",
        "Actualizar demanda, costos de cartón y condiciones logísticas; reoptimizar el catálogo cuando los cambios sean materiales.",
        "Incorporar datos de destino y tarifas intra/extra-zona cuando estén disponibles para refinar el costo logístico.",
        "Usar la carpeta de reproducción entregada para ejecutar el proceso, validar resultados y conservar trazabilidad de cada actualización.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Actualización del catálogo ante nuevos productos", level=1)
    doc.add_paragraph("La alta incremental es la primera alternativa ante un lanzamiento: asignar el nuevo SKU a una de las cajas existentes, verificando capacidad, headspace, resistencia ECT, palletización y costo incremental. Si ninguna opción resulta factible o económicamente conveniente, podrá proponerse un nuevo diseño.")
    doc.add_paragraph("La revisión focalizada evalúa el nuevo SKU y los tipos de caja potencialmente afectados. Permite incorporar productos con menor disrupción operativa, manteniendo inicialmente las demás asignaciones vigentes. Sin embargo, no garantiza el menor costo global: un nuevo SKU puede crear oportunidades de consolidación o modificar las bandas de descuento aplicables por planta.")
    doc.add_paragraph("La reoptimización integral debe ejecutarse ante cambios relevantes en la demanda proyectada, los costos de cartón o las condiciones logísticas. Reevalúa el portafolio completo para identificar la combinación de asignaciones y tipos de caja de menor costo.")
    doc.add_paragraph("La creación de un nuevo tipo debería justificarse cuando ninguna caja existente ofrezca una solución simultáneamente factible y competitiva, o cuando el beneficio económico esperado compense los costos de homologación y transición que Bonsai defina.")
    doc.add_paragraph("La carpeta de código entregada vuelve operativa esta propuesta: implementa la alta incremental, la revisión focalizada, el recálculo con nueva demanda y el flujo de reoptimización integral. El protocolo operativo y los comandos se detallan en la sección 12 del anexo técnico.")

    doc.add_heading("Riesgos y mitigaciones", level=1)
    add_table(doc, ["Riesgo / supuesto", "Mitigación"], [
        ["Dimensiones de 0,1 mm frente a capacidad real de fabricación", "Convertir a tolerancias de plano y validar prototipos; no interpretar como precisión metrológica exigible."],
        ["Demanda histórica usada como proxy anual", "Actualizar operaciones_planta antes de cada ciclo de contratación."],
        ["Flete uniforme de USD 150", "Conservar estructura por región y reemplazar por tarifas/volúmenes reales cuando estén disponibles."],
        ["Reglas geométricas ambiguas en los documentos", "Mantener la interpretación aceptada por el evaluador oficial; documentarla y versionarla."],
    ], [6.0, 10.0])

    doc.add_heading("Beneficios adicionales no cuantificados", level=1)
    doc.add_paragraph("El ahorro informado no incluye los beneficios potenciales asociados a menores costos de inventario, homologación y complejidad administrativa. Si bien la reducción del número de tipos de caja probablemente genere ahorros adicionales en esas áreas, estos no se incorporan a la estimación de USD 21,16 M por falta de información que permita cuantificarlos.")
    doc.add_paragraph("La magnitud del ahorro estimado justifica avanzar con prioridad a la validación industrial. El piloto deberá confirmar los costos de implementación, las tolerancias de fabricación y el desempeño físico antes de decidir la transición completa.")

    path = OUT / "_fuente_informe_ejecutivo.docx"
    doc.save(path)
    return path


def technical_annex() -> Path:
    doc = base_doc("Anexo técnico: datos, diseño y validación", "Soporte metodológico del informe ejecutivo | Julio de 2026")
    add_callout(doc, "Alcance", "Este anexo documenta el análisis exploratorio de los datos originales, la solución validada y las pruebas que respaldan las decisiones metodológicas.")

    doc.add_heading("1. Resultado reproducible", level=1)
    add_table(doc, ["Elemento", "Valor"], [
        ["Archivo base", "baseline/asignacion_0_1mm.csv"],
        ["SHA-256", "a79aff251e1867a95efedd2548af5014df985f47cabfcd3161c7482a410b0d29"],
        ["Costo total", "USD 188.079.384,24"],
        ["Cartón / flete", "USD 26.921.184,24 / USD 161.158.200,00"],
        ["Pallets / tipos", "1.074.388 / 59"],
        ["Espesor global", "3,0 mm"],
        ["Validación oficial", "Archivo validado; ahorro verificado: 10,11%"],
        ["Certificación del MIP", "Óptima; gap 0,0000% en 1.336 diseños candidatos"],
    ], [5.0, 11.0])
    add_source(doc, "baseline/resumen_0_1mm.json, output_gurobi_certify_1h/resumen_decimal.json y validación del evaluador oficial.")
    doc.add_paragraph("La corrida de certificación resolvió el MIP hasta optimalidad: la mejor solución y la mejor cota coincidieron en USD 188.079.384,24. La garantía cubre el conjunto discreto de 1.336 diseños candidatos generado a 0,1 mm. No equivale a una prueba sobre dimensiones continuas que no formen parte de ese conjunto.")
    doc.add_paragraph("Muestra del formato del entregable final:")
    add_table(doc, ["SKU", "Grosor", "Exterior L × A × H (mm)"], [
        ["BR0001", "3,0", "400,0 × 300,0 × 300,0"],
        ["BR0002", "3,0", "400,0 × 300,0 × 271,9"],
        ["BR0003", "3,0", "400,0 × 299,0 × 178,0"],
        ["BR0004", "3,0", "400,0 × 300,0 × 244,2"],
        ["BR0005", "3,0", "400,0 × 300,0 × 221,1"],
    ], [3.0, 3.0, 10.0])

    doc.add_heading("2. Datos, entidades y trazabilidad", level=1)
    doc.add_paragraph("El catálogo aporta 427 productos y su caja actual; especificaciones_cajas aporta geometría y grosor; operaciones_planta aporta la demanda por producto/planta y es la fuente autorizada de demanda; procurement_cajas aporta precios históricos y se recalcula con las cantidades de operaciones. El modelo une cada producto con su caja de referencia y sus volúmenes por planta.")
    add_table(doc, ["Fuente", "Uso en el modelo", "Observación"], [
        ["catalogo_productos.csv", "Producto, caja actual y dimensiones de referencia", "204 tipos actuales identificados"],
        ["especificaciones_cajas.csv", "Dimensiones y grosor histórico", "Se limpió el campo grosor: valor numérico con o sin sufijo “mm”"],
        ["operaciones_planta.csv", "Demanda por SKU y planta", "Fuente de demanda autorizada"],
        ["procurement_cajas.csv", "Precio de cartón; recalculado", "No se usan volúmenes históricos como demanda"],
        ["Consigna + FAQ corregido", "Restricciones y semántica", "El FAQ corregido reemplaza el anterior"],
    ], [4.0, 7.0, 5.0])
    doc.add_paragraph("Calidad de datos: se normalizó el grosor que aparecía como número o con sufijo “mm”. Los campos unitarios y descuentos de Procurement contienen valores `ERROR` y formatos porcentuales heterogéneos; por ello no se usaron como fuente de cálculo. El costo se recalcula aplicando la regla oficial de precio base y bandas de volumen sobre la demanda de Operaciones. La reconciliación muestra además 88 de 204 tipos con volúmenes distintos entre Procurement y Operaciones.")
    add_source(doc, "procurement_cajas.csv, operaciones_planta.csv; criterio de limpieza y recálculo implementado.")

    doc.add_heading("3. Análisis exploratorio del punto de partida", level=1)
    doc.add_paragraph("El EDA se concentra en variables que explican costo y factibilidad: geometría de caja, uso de pallet, distribución de demanda por planta y concentración de volumen por tipo. No se interpreta la categoría comercial como criterio de consolidación porque no garantiza compatibilidad física.")
    doc.add_picture(str(make_chart_dimensions()), width=Inches(6.6))
    doc.add_picture(str(make_chart_utilization()), width=Inches(6.35))
    add_source(doc, "especificaciones_cajas.csv, catalogo_productos.csv y operaciones_planta.csv. Las líneas verticales distinguen el promedio simple y el ponderado por demanda.")

    doc.add_heading("4. Demanda, plantas y concentración de portafolio", level=1)
    doc.add_picture(str(make_chart_plants()), width=Inches(6.6))
    doc.add_picture(str(make_chart_pareto()), width=Inches(6.35))
    doc.add_paragraph("La demanda está distribuida entre cinco plantas y los tipos de caja presentan una concentración de volumen desigual. Esto respalda dos decisiones: el costo debe calcularse por planta y la consolidación debe ser física y económica, no una simplificación indiscriminada del catálogo.")
    add_source(doc, "operaciones_planta.csv y catalogo_productos.csv. El Pareto ordena los 204 tipos actuales por demanda operacional agregada.")

    doc.add_heading("5. Métricas exigidas: utilización por región y tiers", level=1)
    doc.add_picture(str(make_chart_utilization_before_after()), width=Inches(6.45))
    metrics = solution_metrics()
    regional_rows = []
    for label, key in [
        ("Buenos Aires", "buenos_aires"),
        ("Curitiba", "curitiba"),
        ("Santiago", "santiago"),
        ("Monterrey", "monterrey"),
        ("Bakersfield", "bakersfield"),
    ]:
        before = metrics["before"]["u_pallet"][key] * 100
        after = metrics["after"]["u_pallet"][key] * 100
        regional_rows.append([
            label,
            f"{before:.1f}%".replace(".", ","),
            f"{after:.1f}%".replace(".", ","),
            f"+{after - before:.1f} pp".replace(".", ","),
        ])
    add_table(doc, ["Región / planta", "Actual", "Optimizado", "Cambio"], regional_rows, [5.0, 3.5, 3.5, 4.8])
    doc.add_paragraph("La utilización por región / planta se calcula como el volumen anual de cajas dividido por la capacidad anual de los pallets utilizados en esa ubicación. Es, por lo tanto, una medida ponderada por volumen y no un promedio simple de SKU.")
    add_table(doc, ["Métrica", "Actual", "Optimizado", "Cambio"], [
        ["U_caja ponderada", f"{metrics['before']['u_caja']*100:.1f}%", f"{metrics['after']['u_caja']*100:.1f}%", "−3,3 pp"],
        ["U_pallet total", f"{metrics['before']['u_pallet_total']*100:.1f}%", f"{metrics['after']['u_pallet_total']*100:.1f}%", "+11,3 pp"],
        ["Pallets", "1.193.792", "1.074.388", "−119.404"],
    ], [4.2, 3.5, 3.5, 4.8])
    doc.add_picture(str(make_chart_tiers()), width=Inches(6.45))
    doc.add_paragraph("La consolidación reduce las combinaciones diseño–planta activas de 430 a 158 y desplaza volumen hacia bandas de mayor escala. Es el mecanismo comercial que explica parte del ahorro de cartón; la otra parte proviene de la geometría y del espesor global permitido.")
    add_source(doc, "Evaluación independiente de la asignación final y reglas de Procurement por diseño físico y planta.")

    doc.add_heading("6. Formulación mínima", level=1)
    doc.add_paragraph("Para cada producto p se genera un conjunto finito de diseños factibles D(p). La variable binaria x(p,d) indica si se elige el diseño d. Cada producto elige exactamente uno. Un diseño compartido induce un tipo de caja y puede atender varios productos compatibles. El objetivo minimiza:")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Costo total = Σ costo de cartón(p, diseño) + USD 150 × Σ pallets(planta, diseño)")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(NAVY)
    doc.add_paragraph("El número de tipos se reporta como métrica de simplificación; no es un desempate lexicográfico del objetivo económico.")
    add_source(doc, "Consigna; decisiones de interpretación acordadas: USD 150 por pallet, operaciones_planta como demanda y N_tipos como métrica.")

    doc.add_heading("7. Reglas aplicadas", level=1)
    add_table(doc, ["Tema", "Regla aplicada"], [
        ["Espesor", "Un único espesor global; sólo 3,0; 4,5; o 5,0 mm en la solución oficial."],
        ["Dimensiones", "Cada eje puede variar dentro de ±10% respecto de la referencia; se permite reducir uno o dos ejes si volumen y demás restricciones se cumplen."],
        ["Producto", "Las dimensiones internas actuales son referencia física y su volumen se preserva como mínimo."],
        ["Headspace", "Se calcula por eje y se limita por espesor: 6%, 8% o 10%, con tope absoluto de 40 mm."],
        ["Apilado / ECT", "Se usa el perímetro externo completo: 2 × (largo + ancho). La capacidad ECT se compara con la carga de las capas superiores."],
        ["Pallet", "Orientaciones admitidas, cajas por capa y capas completas determinan pallets."],
        ["Un producto por pallet", "Los pallets se calculan separadamente por SKU, planta y diseño; no se mezclan productos."],
        ["Procurement", "El tier se determina por volumen anual acumulado de cada diseño físico en cada planta."],
    ], [4.0, 12.0])
    add_source(doc, "Consigna, secciones de dimensiones, ECT y pallet; FAQ corregido, preguntas 10, 11, 13 y 22; decisiones de proyecto documentadas.")

    doc.add_heading("8. Cómo se hace tratable el espacio de diseño", level=1)
    doc.add_paragraph("Las dimensiones parecen continuas, pero las decisiones económicas cambian sólo en puntos de quiebre: por ejemplo, cuando cambia el número de cajas por capa, el número de capas, el límite de headspace o una cota dimensional. El generador recorre esos puntos relevantes, crea firmas de compatibilidad y elimina diseños dominados. Con 0,1 mm se comprimieron aproximadamente 3.302 millones de puntos de grilla por producto a 246.820 puntos relevantes agregados, que se reducen a aproximadamente 1,3 mil diseños candidatos y sus vínculos producto–diseño.")
    doc.add_paragraph("Luego, un MIP selecciona el diseño de cada producto y consolida productos con firma común. La corrida de certificación de 1.336 diseños candidatos se resolvió con Gurobi hasta optimalidad, con gap 0,0000%. Es una garantía exacta dentro de ese universo discreto; no es una afirmación de optimalidad sobre geometrías continuas no incluidas. El proyecto conserva formulaciones y validadores abiertos para reproducción sin depender de un único motor comercial.")
    add_source(doc, "baseline/resumen_0_1mm.json; src/bonsai/candidates.py, decimal_candidates.py, optimizer.py y solution_validation.py.")

    doc.add_heading("9. Diagnóstico base y lógica económica", level=1)
    doc.add_paragraph("El punto de partida combina un catálogo fragmentado, demanda distribuida entre plantas y reglas físicas estrictas. El modelo no agrupa productos por su categoría comercial: sólo comparte una caja cuando cada producto puede usar exactamente el mismo diseño respetando todas las restricciones. Así, la simplificación del catálogo se alcanza sin convertir la estandarización en un riesgo de factibilidad.")

    doc.add_heading("10. Validación y límites", level=1)
    doc.add_paragraph("Cada archivo candidato se valida independientemente antes de exportar: una fila por producto, espesor único, geometría, volumen, headspace, ECT, palletización, costo y tipos. El archivo final fue además enviado al evaluador oficial. La precisión de 0,1 mm está aceptada por ese evaluador, pero requiere una capa posterior de especificación industrial con tolerancias, herramentales y pruebas antes de implementación física.")
    add_table(doc, ["No asumir", "Tratamiento correcto"], [
        ["Que el histórico sea el plan futuro", "La demanda se refresca desde operaciones antes de cada contratación."],
        ["Que un espesor histórico esté homologado", "Separa evidencia histórica de regla oficial y de factibilidad comercial."],
        ["Que menos tipos sea siempre mejor", "Se mide y se gestiona, pero no desplaza una solución de menor costo."],
        ["Que la forma del producto sea libre", "La evidencia del validador descarta esa interpretación bajo las reglas del caso."],
    ], [6.0, 10.0])

    doc.add_heading("11. Trazabilidad y reproducibilidad", level=1)
    doc.add_paragraph("La solución se conserva como un archivo de asignación con una fila por SKU, un formato de salida fijo y una verificación independiente de geometría, costo y restricciones. Su huella SHA-256 y el detalle de resultados se documentan al inicio de este anexo para permitir su revisión y reproducción.")
    doc.add_paragraph("Se entrega además una carpeta de reproducción documentada. Incluye los datos de entrada que deben proveerse, el flujo principal de optimización, la ruta decimal que genera el CSV final, un validador independiente, el generador de métricas y las dependencias requeridas. El archivo README indica los comandos de ejecución.")

    doc.add_heading("12. Protocolo de actualización del catálogo y nuevos productos", level=1)
    doc.add_paragraph("El protocolo contempla tres modos complementarios. El alta incremental incorpora un nuevo SKU asignándolo a una caja existente cuando cumple capacidad, headspace, resistencia ECT, palletización y costo incremental. Si ninguna alternativa es factible o competitiva, se justifica un nuevo diseño.")
    doc.add_paragraph("La revisión focalizada libera el nuevo SKU y los tipos de caja potencialmente afectados, manteniendo fijas las demás asignaciones. Permite evaluar oportunidades de consolidación y cambios en las bandas de Procurement con un alcance acotado.")
    doc.add_paragraph("La reoptimización integral libera el portafolio completo. Debe utilizarse ante cambios materiales de demanda, costos de cartón o condiciones logísticas, porque busca nuevamente la combinación global de menor costo y puede reasignar productos existentes.")
    doc.add_paragraph("La carpeta de código entregada implementa los tres modos. El alta incremental selecciona el tipo vigente factible de menor costo incremental. La revisión focalizada identifica, por capas de compatibilidad, los tipos vigentes alcanzables desde el nuevo SKU y libera sólo los SKU de ese vecindario; los reasigna entre los tipos físicos activos, mientras conserva fijas las demás asignaciones. El flujo principal permite ejecutar la reoptimización integral. Además, el recálculo con nueva demanda conserva catálogo y asignaciones, reemplaza los volúmenes anuales por una proyección y recalcula pallets, bandas de Procurement, cartón, flete y costo total.")
    doc.add_paragraph("De esta forma, la solución entregada es escalable y reutilizable: puede operar sobre nuevas proyecciones de demanda y futuros lanzamientos, sin reconstruir manualmente el análisis. Los contratos de entrada, comandos y archivos de salida están documentados en el README de la carpeta de reproducción.")
    doc.add_paragraph("Los contratos de entrada, los comandos y los archivos de salida de estos modos están documentados en el README de la carpeta de reproducción entregada.")
    add_source(doc, "Consigna: extensibilidad y protocolo de nuevos productos; decisión de arquitectura acordada en el proyecto.")

    path = OUT / "_fuente_anexo_tecnico.docx"
    doc.save(path)
    return path


if __name__ == "__main__":
    executive_path = executive_report()
    annex_path = technical_annex()
    master = Document(executive_path)
    master.add_page_break()
    composer = Composer(master)
    composer.append(Document(annex_path))
    combined_path = OUT / "Informe_Bonsai_Corp_Final_v17.docx"
    composer.save(combined_path)
    print(combined_path)
