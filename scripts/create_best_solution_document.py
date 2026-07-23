"""Build the concise technical note for the current validated Bonsai solution."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "output_documentation" / "Como_llegamos_a_la_mejor_solucion_actual_tecnicas.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "5E6B78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
WHITE = "FFFFFF"


def set_run_font(run, *, size: float, color: str = "000000", bold: bool = False, italic: bool = False) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, *, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for key, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = margins.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int], *, indent: int = 120) -> None:
    """Apply fixed DXA geometry required by the document design preset."""
    total = sum(widths)
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = tbl.tblGrid
    for grid_col, width in zip(grid.gridCol_lst, widths):
        grid_col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def style_cell(cell, text: str, *, size: float = 10.2, bold: bool = False, color: str = "000000", align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.08
    paragraph.clear()
    run = paragraph.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold)


def add_table(doc, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for cell, text in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, LIGHT_BLUE)
        style_cell(cell, text, size=9.8, bold=True, color=INK)
    for row_values in rows:
        cells = table.add_row().cells
        for index, (cell, text) in enumerate(zip(cells, row_values)):
            if index == 0:
                set_cell_shading(cell, LIGHT_GRAY)
            style_cell(cell, text, size=9.7, bold=index == 0)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_heading(doc, text: str, level: int = 1) -> None:
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.add_run(text)


def add_body(doc, text: str) -> None:
    paragraph = doc.add_paragraph(style="Normal")
    paragraph.add_run(text)


def add_labeled_body(doc, label: str, text: str) -> None:
    paragraph = doc.add_paragraph(style="Normal")
    label_run = paragraph.add_run(f"{label} ")
    set_run_font(label_run, size=11, color=INK, bold=True)
    body_run = paragraph.add_run(text)
    set_run_font(body_run, size=11)


def setup_document() -> Document:
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.paragraph_format.space_after = Pt(0)
    header_run = header.add_run("Bonsai Corp | Optimización de packaging")
    set_run_font(header_run, size=8.5, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_before = Pt(0)
    footer_run = footer.add_run("Nota técnica | Estado de solución validada")
    set_run_font(footer_run, size=8.5, color=MUTED)
    return document


def build_document() -> None:
    doc = setup_document()

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(12)
    title.paragraph_format.space_after = Pt(4)
    title_run = title.add_run("Cómo se llegó a la mejor solución actual")
    set_run_font(title_run, size=23, color=INK, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(16)
    subtitle_run = subtitle.add_run(
        "Bonsai Corp - optimización de packaging y logística | Nota técnica para Sistemas, Datos e IT"
    )
    set_run_font(subtitle_run, size=12.2, color=MUTED)

    metadata = doc.add_table(rows=3, cols=2)
    metadata.style = "Table Grid"
    set_table_geometry(metadata, [2700, 6660])
    for row, (label, value) in zip(
        metadata.rows,
        (
            ("Estado", "Mejor CSV validado localmente"),
            ("Costo total", "USD 188.092.808,10"),
            ("Score estimado", "10,104560% frente al histórico de USD 209.235.093,94"),
        ),
    ):
        set_cell_shading(row.cells[0], LIGHT_BLUE)
        style_cell(row.cells[0], label, size=10.2, bold=True, color=INK)
        style_cell(row.cells[1], value, size=10.2, bold=label == "Costo total")

    add_heading(doc, "1. Resumen ejecutivo")
    add_body(doc,
        "La solución actual se obtuvo conservando una formulación de factibilidad estricta y mejorando de manera incremental una solución exacta de 3 mm. "
        "El costo final validado es USD 188.092.808,10. La mejora acumulada respecto de la primera solución exacta de 3 mm es USD 7.326,06."
    )
    add_body(doc,
        "La idea central fue combinar una base global factible con dos refinamientos de costo: primero, una reoptimización exacta de todos los SKUs que demandan Buenos Aires; después, una relajación lineal global usada únicamente para proponer pocos destinos por SKU y una reparación entera exacta sobre ese pool. "
        "Cada resultado aceptado fue recalculado de forma independiente desde el CSV antes de convertirse en incumbente."
    )

    add_heading(doc, "2. Qué problema se está minimizando")
    add_labeled_body(
        doc,
        "Variable de decisión.",
        "Para cada código de producto se elige una única caja, definida por grosor global, dimensiones exteriores enteras en milímetros y, por derivación, dimensiones internas.",
    )
    add_labeled_body(
        doc,
        "Costo de packaging.",
        "Para cada tipo físico de caja y planta se agregan las unidades anuales y se aplica el precio unitario correspondiente al tramo acumulado de procurement. El descuento es all-units: al alcanzar un tramo, ese precio se aplica a todas las unidades de ese tipo en esa planta.",
    )
    add_labeled_body(
        doc,
        "Costo de flete.",
        "Se calcula el número entero de pallets por SKU y planta a partir de la capacidad de la caja exterior; se valorizan todos los pallets a USD 150, conforme al supuesto aprobado.",
    )
    add_labeled_body(
        doc,
        "Objetivo.",
        "Minimizar packaging + flete. El número de tipos de caja se informa como métrica, pero no se usa como desempate lexicográfico.",
    )

    add_heading(doc, "3. Reglas de factibilidad que permanecen fijas")
    add_labeled_body(
        doc,
        "Demanda.",
        "La fuente de demanda es operaciones_planta.csv. Procurement se recalcula a partir de esas cantidades y de la asignación propuesta.",
    )
    add_labeled_body(
        doc,
        "Grosor y dimensiones.",
        "La mejor solución usa grosor global de 3 mm. Las dimensiones exteriores se redondean a milímetros enteros y se cumple exterior = interior + 2 x grosor por eje.",
    )
    add_labeled_body(
        doc,
        "Producto, volumen y headspace.",
        "Se respeta el volumen físico del producto y las tolerancias documentadas por eje. El headspace puede ubicarse en cualquiera de los tres ejes; no se permite inventar una regla de orientación adicional. La compatibilidad se verifica con el mismo validador independiente que procesa el CSV de salida.",
    )
    add_labeled_body(
        doc,
        "Pallet y ECT.",
        "La capacidad por pallet se calcula con las dimensiones exteriores. El perímetro utilizado para ECT es el externo, según el criterio confirmado.",
    )

    add_heading(doc, "4. Camino que produjo la mejora")
    add_body(doc,
        "La siguiente tabla muestra solamente las etapas que aportaron una reducción de costo. Cada fila parte del CSV validado de la fila anterior."
    )
    add_table(
        doc,
        ["Etapa", "Criterio aplicado", "Costo total (USD)", "Ahorro de la etapa (USD)"],
        [
            ["Base exacta 3 mm", "Modelo exacto sobre el universo de cajas factibles y no dominadas.", "188.100.134,16", "-"],
            ["Refinamiento por destino", "Reasignación coordinada hacia una caja destino rentable, con verificación exacta de procurement y pallets.", "188.098.269,54", "1.864,62"],
            ["SCIP Buenos Aires", "MIP exacto: se liberaron los 134 SKUs con demanda en Buenos Aires y se mantuvo fijo el resto; cada SKU pudo elegir cualquier caja exacta compatible.", "188.092.856,34", "5.413,20"],
            ["Pool guiado por LP", "Relajación lineal global para proponer hasta 4 destinos por SKU, seguida por un MIP entero exacto. Se aceptó sólo la solución con menor costo validado.", "188.092.808,10", "48,24"],
        ],
        [1700, 3600, 2050, 2010],
    )

    add_heading(doc, "5. Técnicas empleadas, en lenguaje práctico")
    add_labeled_body(
        doc,
        "MIP (Mixed-Integer Programming).",
        "Es un modelo matemático de decisiones combinatorias. En este caso, cada SKU debe elegir exactamente una caja; esa elección se representa con variables binarias. El modelo respeta a la vez compatibilidad, pallets, flete y los tramos de procurement, por lo que puede comparar combinaciones de cambios que no serían evidentes SKU por SKU.",
    )
    add_labeled_body(
        doc,
        "SCIP.",
        "Es el motor de optimización que resuelve el MIP. Se usa a través de OR-Tools. SCIP explora combinaciones, mantiene la mejor solución factible conocida (incumbente) y calcula cotas para saber cuánto margen teórico queda. El costo comunicado por SCIP nunca se toma como suficiente: el CSV resultante se vuelve a validar fuera del solver.",
    )
    add_labeled_body(
        doc,
        "LP (Linear Programming) o relajación lineal.",
        "Es la misma formulación, pero permitiendo temporalmente que una asignación sea fraccional, por ejemplo 0,7 de una caja y 0,3 de otra. Ese resultado no es una solución operativa ni se entrega. Su utilidad es analítica: revela cuáles cajas parecen atractivas y qué decisiones tienen mejor señal económica antes de imponer la integridad de una caja por SKU.",
    )
    add_labeled_body(
        doc,
        "Pool guiado por LP.",
        "Es una lista corta de candidatos por SKU. Incluye siempre la caja incumbente y suma las alternativas favorecidas por la relajación lineal, ya sea porque recibieron masa fraccional positiva o porque tienen costo reducido atractivo. Después, el MIP elige de forma entera y exacta dentro de ese pool; el pool guía la búsqueda, no sustituye la validación ni las restricciones.",
    )
    add_labeled_body(
        doc,
        "LNS (Large Neighborhood Search) o vecindario grande.",
        "En vez de reoptimizar todos los SKUs en cada paso, se libera un subconjunto conectado de decisiones y se mantiene fijo el resto. En la mejora por destino se liberaron grupos que podían consolidarse en una caja; en la mejora de Buenos Aires se liberaron todos los SKUs que tenían demanda en esa planta. Esto concentra el esfuerzo de SCIP donde hay interacción económica real.",
    )

    add_heading(doc, "6. Detalle técnico de las etapas ganadoras")
    add_heading(doc, "6.1 Base exacta de 3 mm", level=2)
    add_body(doc,
        "Se generó el universo de alternativas con dimensiones enteras y se descartaron las cajas que no cumplían simultáneamente compatibilidad de producto, reglas de geometría, ECT, palletización o dominancia de costo/capacidad. "
        "La capa de optimización decide una caja por SKU y calcula de forma integrada los tramos de procurement por tipo físico y planta. Esta base aportó una solución factible de costo USD 188.100.134,16."
    )
    add_heading(doc, "6.2 Refinamiento coordinado por destino", level=2)
    add_body(doc,
        "Sobre la base exacta, se evaluaron movimientos coordinados de grupos completos de SKUs hacia un mismo diseño destino. El valor de un movimiento no se calculó con un precio unitario fijo: se recalcularon simultáneamente las unidades acumuladas del tipo origen y destino, el tramo de procurement y los pallets. "
        "Esto permitió encontrar una combinación que redujo el costo total USD 1.864,62 y redujo además el número de tipos de 55 a 53."
    )
    add_heading(doc, "6.3 MIP exacto focalizado en Buenos Aires", level=2)
    add_body(doc,
        "El siguiente salto se logró con SCIP. Se fijaron las asignaciones de los SKUs sin demanda en Buenos Aires como constantes de costo y volumen. Los 134 SKUs restantes conservaron acceso a todo el universo exacto de 1.248 diseños compatibles. "
        "Así, el modelo pudo intercambiar packaging por flete cuando la reducción de pallets compensaba el cambio de procurement. El resultado modificó 19 SKUs, redujo 84 pallets y produjo un ahorro neto de USD 5.413,20."
    )
    add_heading(doc, "6.4 Pool de candidatos guiado por relajación lineal", level=2)
    add_body(doc,
        "Para el último refinamiento se resolvió la relajación lineal de la formulación global. Las asignaciones fraccionales y los costos reducidos se usaron como señal para construir un conjunto pequeño de candidatos por SKU, siempre incluyendo la caja incumbente. "
        "Luego se resolvió el modelo entero con SCIP sobre ese pool de cuatro opciones por SKU. Esta etapa preservó los pallets y el flete de la etapa anterior, redujo packaging en USD 48,24 y disminuyó los tipos físicos de 53 a 51."
    )

    add_heading(doc, "7. Resultado actual y controles")
    add_table(
        doc,
        ["Métrica", "Valor validado"],
        [
            ["Packaging", "USD 26.892.608,10"],
            ["Flete", "USD 161.200.200,00"],
            ["Costo total", "USD 188.092.808,10"],
            ["Pallets", "1.074.668"],
            ["Tipos físicos", "51"],
            ["Score estimado", "10,104560%"],
        ],
        [3000, 6360],
    )
    add_labeled_body(
        doc,
        "Archivo de entrega.",
        "output_lp_pool_after_ba_15m/asignacion_optima.csv",
    )
    add_labeled_body(
        doc,
        "Huella SHA-256.",
        "3050586EE9C111D417C3587F62F8D4CEE80DDA8E6419F34E12B9744377740DEF",
    )
    add_labeled_body(
        doc,
        "Validación.",
        "El CSV fue leído nuevamente por solution_validation.validate_solution_csv, que verifica columnas, unicidad de SKU, grosor global, dimensiones enteras, compatibilidad geométrica y recálculo completo de packaging, pallets y flete.",
    )

    add_heading(doc, "8. Reproducibilidad para Sistemas y Datos")
    add_body(doc,
        "La solución se puede reconstruir desde los CSV fuente y el código Python del repositorio. La responsabilidad de cada componente está separada: data.py carga y normaliza; exact_candidates.py genera el universo discreto; scip_optimizer.py formula el MIP; lp_pool.py construye los pools guiados por la relajación; costs.py calcula el costo; y solution_validation.py vuelve a validar el archivo de salida."
    )
    add_table(
        doc,
        ["Componente", "Responsabilidad"],
        [
            ["operaciones_planta.csv", "Demanda anual por SKU y planta."],
            ["especificaciones_cajas.csv", "Parámetros de cartón y grosor, normalizados antes del modelado."],
            ["src/bonsai/exact_candidates.py", "Generación de cajas factibles en milímetros enteros."],
            ["src/bonsai/scip_optimizer.py", "Optimización entera con procurement y flete integrados."],
            ["src/bonsai/lp_pool.py", "Selección de candidatos guiada por la relajación lineal."],
            ["src/bonsai/solution_validation.py", "Validación independiente del CSV final y recálculo de costo."],
        ],
        [3300, 6060],
    )
    add_body(doc,
        "Principio operativo: una mejora no se considera válida por el objetivo que informa el solver. Se escribe el CSV, se relee sin usar el estado interno del solver y se acepta sólo si el costo recalculado es estrictamente menor que el incumbente."
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.core_properties.title = "Cómo se llegó a la mejor solución actual - Bonsai Corp"
    doc.core_properties.subject = "Metodología y reproducibilidad de la solución de packaging"
    doc.core_properties.author = "Equipo de optimización"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
